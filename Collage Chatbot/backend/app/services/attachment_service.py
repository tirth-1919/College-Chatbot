"Owned student attachment storage and deterministic text extraction."""
import csv
import hashlib
import io
import json
import logging
import os
import shutil
from pathlib import Path
from typing import Any, Dict, Tuple
from openpyxl import load_workbook
from PyPDF2 import PdfReader
from docx import Document
from backend.app.models.entities import Attachment
from backend.app.security.file_validator import FileSecurityValidator
logger = logging.getLogger(__name__)
SAFE_ERROR = "I couldn't read this file. Please try another copy."

class AttachmentService:
    def __init__(self, root: str = "./storage/attachments"):
        self.root = Path(root).resolve()
        self.root.mkdir(parents=True, exist_ok=True)
        self.validator = FileSecurityValidator()

    def _extract(self, content: bytes, suffix: str) -> Tuple[str, Dict[str, Any]]:
        if suffix == ".pdf":
            reader = PdfReader(io.BytesIO(content))
            pages = [{"page": i + 1, "text": (p.extract_text() or "").strip()} for i, p in enumerate(reader.pages)]
            return "\n\n".join(f"Page {p['page']}\n{p['text']}" for p in pages), {"pages": len(pages), "page_text": pages}
        if suffix == ".docx":
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables = [[[cell.text.strip() for cell in row.cells] for row in table.rows] for table in doc.tables]
            return "\n".join(paragraphs + [" | ".join(row) for table in tables for row in table]), {"paragraphs": len(paragraphs), "tables": tables}
        if suffix in {".txt", ".md"}:
            return content.decode("utf-8", errors="strict"), {}
        if suffix == ".csv":
            rows = list(csv.reader(io.StringIO(content.decode("utf-8", errors="strict"))))
            return "\n".join(" | ".join(row) for row in rows), {"columns": rows[0] if rows else [], "rows": max(len(rows) - 1, 0)}
        if suffix == ".xlsx":
            book = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            sheets = {}
            lines = []
            for sheet in book.worksheets:
                rows = [[("" if value is None else str(value)) for value in row] for row in sheet.iter_rows(values_only=True)]
                sheets[sheet.title] = {"rows": len(rows), "columns": max((len(row) for row in rows), default=0)}
                lines.append(f"Sheet: {sheet.title}\n" + "\n".join(" | ".join(row) for row in rows))
            return "\n\n".join(lines), {"sheets": sheets}
        if suffix in {".jpg", ".jpeg", ".png", ".gif", ".webp"}:
            return "", {"image": True}
        raise ValueError("unsupported format")

    def save(self, file_name: str, content: bytes, file_type: str) -> Tuple[str, str, Dict[str, Any]]:
        suffix = Path(file_name).suffix.lower()
        digest = hashlib.sha256(content).hexdigest()
        extracted, metadata = self._extract(content, suffix)
        path = self.root / digest[:2] / f"{digest}{suffix}"
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(content)
        return str(path), digest, {**metadata, "extracted_text": extracted}

    def delete_storage(self, path: str) -> None:
        candidate = Path(path).resolve()
        if candidate != self.root and self.root not in candidate.parents:
            logger.warning("Refusing to delete attachment outside storage root")
            return
        try:
            candidate.unlink(missing_ok=True)
            if candidate.parent != self.root and not any(candidate.parent.iterdir()):
                candidate.parent.rmdir()
        except OSError:
            logger.exception("Attachment cleanup failed")
