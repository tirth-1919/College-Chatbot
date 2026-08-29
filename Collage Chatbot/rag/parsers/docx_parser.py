import io
import hashlib
from typing import Dict, Any, Optional, List

class DOCXParser:
    """
    DOCX document parser for extracting text, metadata, and structure from Word documents.
    Supports basic text extraction, paragraph-level processing, and metadata extraction.
    """
    
    def __init__(self):
        self.supported_formats = ['docx']
        self.docx_available = False
        
        try:
            from docx import Document
            self.Document = Document
            self.docx_available = True
        except ImportError:
            print("[DOCXParser] python-docx not installed. Install with: pip install python-docx")
    
    def parse_docx(self, file_bytes: bytes, filename: str = "") -> Dict[str, Any]:
        """
        Parse DOCX file and extract text, metadata, and structure.
        
        Args:
            file_bytes: Raw bytes of the DOCX file
            filename: Original filename of the DOCX
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        if not self.docx_available:
            return {
                "success": False,
                "error": "python-docx not installed",
                "filename": filename
            }
        
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = self.Document(doc_file)
            
            # Extract basic metadata
            metadata = self._extract_metadata(doc, filename)
            
            # Extract text from all paragraphs
            paragraphs_content = []
            full_text = ""
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_content = {
                        "paragraph_number": para_num + 1,
                        "text": paragraph.text.strip(),
                        "style": paragraph.style.name if paragraph.style else "Normal",
                        "char_count": len(paragraph.text)
                    }
                    paragraphs_content.append(para_content)
                    full_text += paragraph.text + "\n"
            
            # Extract tables if present
            tables_content = []
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append({
                    "table_number": table_num + 1,
                    "rows": len(table.rows),
                    "columns": len(table.columns) if table.rows else 0,
                    "data": table_data
                })
            
            # Clean and structure the extracted text
            clean_text = self._clean_text(full_text)
            
            # Generate content hash
            content_hash = hashlib.sha256(clean_text.encode('utf-8')).hexdigest()
            
            # Extract sections if possible
            sections = self._extract_sections(doc.paragraphs)
            
            return {
                "success": True,
                "filename": filename,
                "format": "docx",
                "metadata": metadata,
                "full_text": clean_text,
                "content_hash": content_hash,
                "paragraphs": paragraphs_content,
                "total_paragraphs": len(paragraphs_content),
                "tables": tables_content,
                "total_tables": len(tables_content),
                "total_characters": len(clean_text),
                "sections": sections,
                "word_count": len(clean_text.split())
            }
            
        except Exception as e:
            print(f"[DOCXParser] Error parsing DOCX: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename
            }
    
    def _extract_metadata(self, doc, filename: str) -> Dict[str, Any]:
        """Extract metadata from DOCX document"""
        metadata = {
            "filename": filename,
            "title": "",
            "author": "",
            "subject": "",
            "creator": "",
            "created": "",
            "modified": "",
            "last_modified_by": "",
            "revision": "",
            "paragraph_count": len(doc.paragraphs)
        }
        
        try:
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "creator": core_props.creator or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                    "last_modified_by": core_props.last_modified_by or "",
                    "revision": str(core_props.revision) if core_props.revision else ""
                })
        except Exception as e:
            print(f"[DOCXParser] Error extracting metadata: {e}")
        
        return metadata
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove multiple consecutive newlines
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        
        return text.strip()
    
    def _extract_sections(self, paragraphs) -> List[Dict[str, Any]]:
        """
        Extract sections from paragraphs based on heading styles.
        """
        sections = []
        current_section = {"title": "Introduction", "content": "", "start_para": 0}
        
        for para_num, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check if this paragraph is a heading (based on style)
            style_name = paragraph.style.name if paragraph.style else "Normal"
            is_heading = any(
                style_name.startswith(prefix) 
                for prefix in ['Heading', 'Title', 'Subtitle']
            )
            
            if is_heading:
                # Save previous section
                if current_section["content"].strip():
                    sections.append(current_section.copy())
                
                # Start new section
                current_section = {
                    "title": text,
                    "content": "",
                    "start_para": para_num,
                    "style": style_name
                }
            else:
                current_section["content"] += text + " "
        
        # Add the last section
        if current_section["content"].strip():
            sections.append(current_section)
        
        # If no sections were found, treat entire text as one section
        if not sections:
            sections = [{
                "title": "Document",
                "content": self._clean_text(' '.join(p.text for p in paragraphs)),
                "start_para": 0
            }]
        
        return sections
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if the file format is supported"""
        return filename.lower().endswith('.docx')
    
    def validate_docx(self, file_bytes: bytes):
        """
        Validate if the file is a valid DOCX.
        
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not self.docx_available:
            return False, "python-docx not installed"
        
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = self.Document(doc_file)
            
            # Try to read the first paragraph to validate
            if len(doc.paragraphs) > 0:
                first_para = doc.paragraphs[0]
                _ = first_para.text
                
            return True, ""
            
        except Exception as e:
            return False, f"Invalid DOCX file: {str(e)}"
